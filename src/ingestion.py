"""
ingestion.py — Document loading and text extraction.

Supports: PDF, TXT, DOCX
Stores page_boundaries for PDFs so chunking can assign accurate page numbers.
TXT and DOCX have no real pages — chunking assigns page_hint=1 for all their chunks.
"""

from __future__ import annotations

import re
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple


# ─── Data Model ───────────────────────────────────────────────────────────────

@dataclass
class RawDocument:
    """A single document loaded from disk, before chunking."""
    source:    str
    full_path: str
    text:      str
    num_pages: int = 0
    metadata:  dict = field(default_factory=dict)
    page_boundaries: List[Tuple[int, int]] = field(default_factory=list)

    @property
    def doc_id(self) -> str:
        return hashlib.md5(self.full_path.encode()).hexdigest()[:8]

    def get_page_for_char(self, char_pos: int) -> int:
        """Return the actual page number for a given character position."""
        if not self.page_boundaries:
            return 1
        page = 1
        for char_start, page_num in self.page_boundaries:
            if char_pos >= char_start:
                page = page_num
            else:
                break
        return page


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Remove repeated headers/footers, lone page numbers, and double blank lines."""
    from collections import Counter
    lines   = text.splitlines()
    freq    = Counter(ln.strip() for ln in lines if ln.strip())
    cleaned = []
    for ln in lines:
        stripped = ln.strip()
        if not stripped:
            if cleaned and cleaned[-1] == "":
                continue
            cleaned.append("")
            continue
        if re.fullmatch(r"\d{1,4}", stripped):
            continue
        if freq[stripped] >= 3 and len(stripped) < 80:
            continue
        cleaned.append(ln)
    return "\n".join(cleaned).strip()


# ─── Loaders ──────────────────────────────────────────────────────────────────

def load_pdf(path: Path) -> RawDocument:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required: pip install pypdf")

    reader           = PdfReader(str(path))
    pages_text:      List[str]            = []
    page_boundaries: List[Tuple[int,int]] = []

    char_offset = 0
    for page_num, page in enumerate(reader.pages, start=1):
        page_text = _clean_text(page.extract_text() or "")
        page_boundaries.append((char_offset, page_num))
        pages_text.append(page_text)
        char_offset += len(page_text) + 2   # +2 for "\n\n" separator

    full_text = "\n\n".join(pages_text)

    return RawDocument(
        source=path.name,
        full_path=str(path),
        text=full_text,
        num_pages=len(reader.pages),
        metadata={"file_type": "pdf"},
        page_boundaries=page_boundaries,
    )


def load_txt(path: Path) -> RawDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    text = _clean_text(text)
    return RawDocument(
        source=path.name,
        full_path=str(path),
        text=text,
        num_pages=0,
        metadata={"file_type": "txt"},
        # No page_boundaries — chunking assigns page_hint=1 for all chunks
    )


def load_docx(path: Path) -> RawDocument:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc        = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text       = "\n\n".join(paragraphs)
    text       = _clean_text(text)
    return RawDocument(
        source=path.name,
        full_path=str(path),
        text=text,
        num_pages=0,
        metadata={"file_type": "docx"},
        # No page_boundaries — chunking assigns page_hint=1 for all chunks
    )


# ─── Public API ───────────────────────────────────────────────────────────────

LOADERS = {
    ".pdf":  load_pdf,
    ".txt":  load_txt,
    ".docx": load_docx,
}


def load_documents(data_dir: str | Path) -> List[RawDocument]:
    """Load all supported documents from *data_dir*."""
    data_dir = Path(data_dir)
    if not data_dir.exists():
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    docs: List[RawDocument] = []
    for path in sorted(data_dir.iterdir()):
        suffix = path.suffix.lower()
        if suffix not in LOADERS:
            continue
        print(f"  📄 Loading {path.name} …")
        try:
            doc = LOADERS[suffix](path)
            if len(doc.text.strip()) < 50:
                print(f"     ⚠️  Skipping {path.name} — too little text extracted.")
                continue
            docs.append(doc)
            print(f"     ✅ {len(doc.text):,} chars, {doc.num_pages} pages")
        except Exception as exc:
            print(f"     ❌ Failed to load {path.name}: {exc}")

    if not docs:
        raise ValueError(f"No valid documents found in '{data_dir}'.")
    return docs